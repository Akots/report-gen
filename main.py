import os
import sys
import pandas as pd
import re
from docx import Document
from datetime import datetime, timedelta
from pymorphy2 import MorphAnalyzer
import openpyxl
import copy

from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

from manual_declensions_ua import manual_declensions_ua

DICT_PATH = os.path.join(os.path.dirname(__file__), 'pymorphy2_dicts_uk/data')
morph = MorphAnalyzer(lang='uk', path=DICT_PATH)

def resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        base_path = os.path.dirname(sys.executable)
    else:
        base_path = os.path.dirname(__file__)
    return os.path.join(base_path, relative_path)

EXCEL_FILE       = resource_path('Дані подача РАПОРТІВ.xlsx')
SDD_FILE         = resource_path('1МБ - СДД.xlsx')
OUTPUT_DIR       = resource_path('Результати')

TEMPLATE_SWITCH = {
    '156': {
        'dovidka': resource_path('ШАБЛОНИ/доповідь - Шаблон.docx'),
        'raport': resource_path('ШАБЛОНИ/РАПОРТ СЗЧ - ШАБЛОН.docx'),
    },
    '71': {
        'raport': resource_path('ШАБЛОНИ/РАПОРТ СЗЧ - ШАБЛОН 71а.docx'),
    },
    '47': {
        'dovidka': resource_path('ШАБЛОНИ/доповідь - ШАБЛОН 47.docx'),
    },
}

STYLE_PROFILE = 'new'  # 'new' (дефолт) або 'old'

def set_format_style(profile: str):
    """Встановити стиль форматування для поточного кроку генерації."""
    global STYLE_PROFILE
    STYLE_PROFILE = 'old' if str(profile).lower() == 'old' else 'new'

def get_style() -> str:
    return STYLE_PROFILE

VAR_TO_COL = {
    'ДАННІ_ПРИЗВ':      'ДАННІ_ПРИЗВ',
    'ДАННІ_ІНН':        'ДАННІ_ІНН',
    'ДАННІ_НР':         'ДАННІ_НР',
    'ДАННІ_МПРОЖИВ':    'ДАННІ_МПРОЖИВ',
    'ДАННІ_ТЕЛ':        'ДАННІ_ТЕЛ',
    'ДАННІ_ОСВ':        'ДАННІ_ОСВ',
    'ДАННІ_ВІРА':       'ДАННІ_ВІРА',
    'ДАННІ_РОДИЧ':      'ДАННІ_РОДИЧ',
    'ДАННІ_СІМЕЙСТАН':  'ДАННІ_СІМЕЙСТАН',
    'ДАННІ_НАЦІОН':     'ДАННІ_НАЦІОН',
    'ДАННІ_ВІЙСЬКОВИЙ': 'ДАННІ_ВІЙСЬКОВИЙ',
    'ДАННІ_ПАСПОРТ':    'ДАННІ_ПАСПОРТ',
}

MONTHS = ['січня','лютого','березня','квітня','травня','червня',
          'липня','серпня','вересня','жовтня','листопада','грудня']

def get_investigation_deadline(date, days=10):
    deadline_date = date + timedelta(days=days)
    return deadline_date.strftime("%d.%m.%Y")

def decline_text(text, case):
    text_lower = text.lower().strip()
    if text_lower in manual_declensions_ua and case in manual_declensions_ua[text_lower]:
        return manual_declensions_ua[text_lower][case]
    declined = []
    for word in text.split():
        parsed = morph.parse(word)[0]
        inflected = parsed.inflect({case})
        declined.append(inflected.word if inflected else word)
    return ' '.join(declined)

def normalize_department(s: str) -> str:
    if not s:
        return s
    s = re.sub(r"\s{2,}", " ", s).strip()

    TAIL = r"\b\d+\s+\S+\s+батальйону(?:\s+військової\s+частини\s+[AА]\d+)?\b"

    # Якщо однакові хвости йдуть підряд — згорнути до одного
    s = re.sub(rf"({TAIL})(?:\s+{TAIL})+", r"\1", s, flags=re.IGNORECASE)

    # Якщо хвіст трапився кілька разів (рознесено) — залишити перший
    spans = [m.span() for m in re.finditer(TAIL, s, flags=re.IGNORECASE)]
    if len(spans) > 1:
        for start, end in reversed(spans[1:]):
            s = (s[:start].rstrip() + " " + s[end:]).strip()
        s = re.sub(r"\s{2,}", " ", s)

    return s

def extract_department_fields(title: str):
    """
    Витягує підрозділ із тексту посади, напр.:
      - '... 1 механізованої роти 1 механізованого батальйону (в/ч Аxxxx)'
      - 'Група зв'язку ... штабу 1 механізованого батальйону'
      - 'мінометної батареї 1 механізованого батальйону'
    Повертає (department_nom, department_gen).
    """
    t = (title or "").strip().replace("’", "'")

    patterns = [
        # 1) ГРУПА штабу ... батальйону [в/ч]
        r"(Група\s+[^,\n]*?\s+штабу\s+\d+\s+\S+\s+батальйону(?:\s+військової\s+частини\s+[AА]\d+)?)",
        # 2) ... роти ... батальйону [в/ч]
        r"(\d+\s+\S+\s+роти\s+\d+\s+\S+\s+батальйону(?:\s+військової\s+частини\s+[AА]\d+)?)",
        # 3) (НОВЕ) *батареї* ... батальйону [в/ч]
        #    захоплюємо прикметник перед 'батареї' (напр. 'мінометної батареї')
        r"((?:\S+\s+)?батареї\s+\d+\s+\S+\s+батальйону(?:\s+військової\s+частини\s+[AА]\d+)?)",
        # 4) штабу ... батальйону [в/ч]
        r"(штабу\s+\d+\s+\S+\s+батальйону(?:\s+військової\s+частини\s+[AА]\d+)?)",
        # 5) ... батальйону [в/ч]
        r"(\d+\s+\S+\s+батальйону(?:\s+військової\s+частини\s+[AА]\d+)?)",
    ]

    department_nom = None
    for pat in patterns:
        m = re.search(pat, t, flags=re.IGNORECASE)
        if m:
            department_nom = m.group(1).strip()
            break

    if not department_nom:
        # Фолбек: беремо від останнього "якоря" до кінця.
        # ВАЖЛИВО: 'батареї' стоїть ПЕРЕД 'взводу/відділення', щоб віддати перевагу рівню батареї.
        anchor_pat = r"(батареї|роти|штабу|батальйону|взводу|відділення)(.*)$"
        m2 = list(re.finditer(anchor_pat, t, flags=re.IGNORECASE))
        if m2:
            department_nom = (m2[-1].group(0)).strip()
        else:
            # останній запасний варіант — права частина після тире
            parts = re.split(r"\s*[—–\-]\s*", t)
            department_nom = parts[-1].strip() if parts else t

    # Нормалізуємо (прибираємо дубль хвоста, зайві пробіли тощо)
    department_nom = normalize_department(department_nom)

    department_gen = decline_text(department_nom, "gent")
    return department_nom, department_gen

def is_empty(val):
    if val is None:
        return True
    if pd.isna(val):
        return True
    sval = str(val).strip()
    if sval == '' or sval.lower() == 'nan':
        return True
    return False

def date_parts(dt):
    return {
        'СЗЧ_ДАТА':      dt.strftime('%d.%m.%Y'),
        'ДОПОВІДЬ_ДАТА': dt.strftime('%d.%m.%Y'),
        'ДОПОВІДЬ_ДЕНЬ': str(dt.day),
        'ДОПОВІДЬ_МІСЯЦЬ':MONTHS[dt.month-1],
        'ДОПОВІДЬ_РІК':  str(dt.year),
        'РАПОРТ_ДАТА':   dt.strftime('%d.%m.%Y'),
        'РАПОРТ_ДЕНЬ':   str(dt.day),
        'РАПОРТ_МІСЯЦЬ': MONTHS[dt.month-1],
        'РАПОРТ_РІК':    str(dt.year),
    }

def format_date(val):
    if pd.notna(val):
        return pd.to_datetime(val, dayfirst=True).strftime('%d.%m.%Y')
    return ''

def format_phone(val):
    if pd.isna(val):
        return ''
    digits = re.sub(r'\D', '', str(val))
    if len(digits) == 10:
        return f"({digits[:3]}) {digits[3:6]}-{digits[6:8]}-{digits[8:10]}"
    return str(val)

def format_inn(val):
    if pd.isna(val):
        return ''
    try:
        return str(int(float(val)))
    except (ValueError, TypeError):
        return str(val)

def default_format(val):
    if pd.notna(val) and str(val).upper() != 'N/A':
        return str(val)
    return ''

formatters = {
    'ДАННІ_НР':  format_date,
    'ДАННІ_ТЕЛ': format_phone,
    'ДАННІ_ІНН': format_inn,
}

def to_genitive(text):
    if not isinstance(text, str): return ''
    words = text.split()
    inflected = [morph.parse(w)[0].inflect({'gent'}).word if morph.parse(w)[0].inflect({'gent'}) else w for w in words]
    return ' '.join(inflected)

def cap_ua(s: str) -> str:
    if not isinstance(s, str): return ''
    # Іван-Григорович -> Іван-Григорович
    return '-'.join(w[:1].upper() + w[1:].lower() if w else '' for w in s.split('-'))

def format_name(full: str, style: str = None) -> str:
    """ПІБ у залежності від стилю:
       new: 'Власюк Іван Іванович'
       old: 'ВЛАСЮК Іван Іванович' (як було раніше)
    """
    st = (style or get_style())
    if not isinstance(full, str): return ''
    parts = full.split()
    if not parts: return ''
    if st == 'old':
        surname = parts[0].upper()
    else:
        surname = cap_ua(parts[0])
    rest = [cap_ua(w) for w in parts[1:]]
    return ' '.join([surname] + rest)

def short_fio(full: str, style: str = None) -> str:
    """Коротка форма: 'І. Власюк' (new) або 'І. ВЛАСЮК' (old)"""
    st = (style or get_style())
    if not full or not isinstance(full, str): return ""
    parts = full.strip().split()
    if len(parts) < 2:
        return full.upper() if st == 'old' else cap_ua(full)
    if st == 'old':
        surname = parts[0].upper()
    else:
        surname = cap_ua(parts[0])
    name_initial = parts[1][:1].upper() + '.'
    return f"{name_initial} {surname}"


def replace_placeholders(doc, ctx):
    for p in doc.paragraphs:
        for key, val in ctx.items():
            if f'{{{{{key}}}}}' in p.text and val is not None:
                for run in p.runs:
                    run.text = run.text.replace(f'{{{{{key}}}}}', val)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                replace_placeholders(cell, ctx)

def replace_all_placeholders_xml(doc, data):
    def iter_block_items(parent):
        for child in parent.element.body:
            if isinstance(child, CT_P):
                yield child
            elif isinstance(child, CT_Tbl):
                yield child
    for block in iter_block_items(doc):
        for node in block.iter():
            if node.tag.endswith('}t') and node.text:
                for key, val in data.items():
                    if key in node.text:
                        node.text = node.text.replace(key, val)
#Получити під шаблон по назві тега
def extract_block_by_tags(doc, tag_name):
    start_tag = f"{{{{{tag_name}}}}}"
    end_tag = f"{{{{/{tag_name}}}}}"
    blocks = []
    in_block = False
    for p in doc.paragraphs:
        if start_tag in p.text:
            in_block = True
            continue  # не включати тег
        elif end_tag in p.text:
            in_block = False
            break
        elif in_block:
            blocks.append(copy.deepcopy(p._element))
    return blocks

def short_fio(full_name):
    if not full_name or not isinstance(full_name, str): return ""
    parts = full_name.strip().split()
    if len(parts) < 2: return full_name
    surname = parts[0]
    name_initial = parts[1][0] + '.' if len(parts[1]) > 0 else ''
    return f"{name_initial} {surname}"

def update_excel_cells(excel_path, fio_to_data, fio_col_name='ПІБ'):
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active  # або ws = wb["Назва_листа"]
    headers = [cell.value for cell in ws[1]]
    col_idx = {name: idx for idx, name in enumerate(headers)}
    fio_map = {fio.strip().lower(): fio for fio in fio_to_data.keys()}
    for row in ws.iter_rows(min_row=2):
        fio_cell = row[col_idx[fio_col_name]].value
        if fio_cell is None:
            continue
        fio_val = str(fio_cell).strip().lower()
        if not fio_val:
            continue
        if fio_val in fio_map:
            true_fio = fio_map[fio_val]
            for k, v in fio_to_data[true_fio].items():
                if k in col_idx and v not in [None, '', 'nan']:
                    # Форматування для ДАННІ_НР
                    if k == 'ДАННІ_НР':
                        if isinstance(v, (pd.Timestamp, datetime)):
                            v = v.strftime('%d.%m.%Y')
                        elif isinstance(v, str):
                            try:
                                v = pd.to_datetime(v, dayfirst=True).strftime('%d.%m.%Y')
                            except:
                                pass
                    row[col_idx[k]].value = v
    wb.save(excel_path)

def insert_block_by_placeholder(doc, placeholder, block_elements):
    # Глибокий пошук у всіх параграфах і всіх таблицях (у параграфах і cell)
    found = False
    def search_and_replace(paragraphs):
        nonlocal found
        for p in paragraphs:
            full_text = ''.join(run.text for run in p.runs)
            if placeholder in full_text:
                parent = p._element.getparent()
                idx = list(parent).index(p._element)
                parent.remove(p._element)
                for el in reversed(block_elements):
                    parent.insert(idx, copy.deepcopy(el if hasattr(el, 'tag') else el._element))
                found = True
                return True
        return False
    if search_and_replace(doc.paragraphs):
        return
    def search_in_tables(tables):
        nonlocal found
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    if search_and_replace(cell.paragraphs):
                        return True
                    if search_in_tables(cell.tables):
                        return True
        return False
    search_in_tables(doc.tables)
    if not found:
        print(f"[insert_block_by_placeholder] ⚠️ Не знайдено плейсхолдер '{placeholder}' у документі")

def _time_from_any(val):
    if val is None: return None
    if isinstance(val, (pd.Timestamp, datetime)):
        return val.hour, val.minute
    # Excel дробова частка доби
    if isinstance(val, (int, float)) and not pd.isna(val):
        total_seconds = int(round(float(val) * 24 * 3600))
        return (total_seconds // 3600) % 24, (total_seconds % 3600) // 60
    s = str(val).strip()
    m = re.search(r'(\d{1,2})[:\.](\d{1,2})', s)
    if m:
        h, mi = int(m.group(1)), int(m.group(2))
        return max(0, min(23, h)), max(0, min(59, mi))
    try:
        dt = pd.to_datetime(s, errors='coerce', dayfirst=True)
        if not pd.isna(dt):
            return dt.hour, dt.minute
    except Exception:
        pass
    return None

def format_time_str(val, style: str = None) -> str:
    """Час у залежності від стилю:
       new: 07.10
       old: 07:10
    """
    st = (style or get_style())
    hm = _time_from_any(val)
    if not hm: return ''
    h, m = hm
    if st == 'old':
        return f"{h:02d}:{m:02d}"
    else:
        return f"{h:02d}.{m:02d}"

# === 1. Читання даних ===
os.makedirs(OUTPUT_DIR, exist_ok=True)
df = pd.read_excel(EXCEL_FILE)
df_sdd = pd.read_excel(SDD_FILE)
df_sdd.columns = [col.strip() for col in df_sdd.columns]
SDD_FIO_COL = 'Прізвище, імя , по батькові'

SDD_MAP = {
    'ПІБ': 'Прізвище, імя , по батькові',
    'ДАННІ_ІНН': 'Ідентифікаційний код',
    'ДАННІ_НР': 'Дата народження',
    'ДАННІ_ПРИЗВ': 'Коли і яким ВК призваний',
    'ДАННІ_МПРОЖИВ': 'Фактичне місце проживання',
    'ДАННІ_ТЕЛ': 'Номер мобільного',
    'ДАННІ_ОСВ': 'Освіта',
    'ДАННІ_РОДИЧ': 'Близькі родичі',
    'ДАННІ_СІМЕЙСТАН': 'Сімейний стан',
    'ДАННІ_НАЦІОН': 'Національність',
    'ДАННІ_ПАСПОРТ': 'Серія, № паспорта',
    'ДАННІ_ВІЙСЬКОВИЙ': '№ військового квитка (посвідчення офіцера)',
}

# === 2. Підготовка нових даних для запису ===
fio_to_data = {}
for idx, row in df.iterrows():
    if not pd.isna(row.get('ДАННІ_ІНН')) and str(row.get('ДАННІ_ІНН')).strip() != '':
        continue
    fio = str(row.get('ПІБ', '')).strip()
    if not fio:
        continue
    sdd_match = df_sdd[df_sdd[SDD_FIO_COL].astype(str).str.strip() == fio]
    if not sdd_match.empty:
        sdd_row = sdd_match.iloc[0]
        newdata = {}
        for k, sdd_col in SDD_MAP.items():
            if k == 'ПІБ':
                continue
            val = sdd_row.get(sdd_col, None)
            if pd.notna(val) and str(val).strip() != '':
                newdata[k] = val
        if newdata:
            fio_to_data[fio] = newdata

if fio_to_data:
    update_excel_cells(EXCEL_FILE, fio_to_data, fio_col_name='ПІБ')
    df = pd.read_excel(EXCEL_FILE)

# === 4. Генерація документів ===

for _, row in df.iterrows():
    format_value = str(row.get('ФОРМАТ', '')).strip()
    maino_value = str(row.get('МАЙНО', '')).strip()
    templates = TEMPLATE_SWITCH.get(format_value, {})

    if not templates:
        print(f"[УВАГА] Формат '{format_value}' не знайдено у TEMPLATE_SWITCH — пропущено.")
        continue

    fio = row.get('ПІБ', '')
    if not isinstance(fio, str) or not fio.strip():
        print("[пропущено] Порожнє або некоректне ПІБ")
        continue

    set_format_style('old' if format_value == "47" else 'new')

    ctx = {}

    base_posada = str(row.get('ПОСАДА', '') or '')
    department, _ = extract_department_fields(base_posada)
    ctx['ПОСАДА'] = base_posada
    ctx['ПОСАДА_РОД'] = to_genitive(base_posada)

    ctx['КОМУ_КЛОПОЧЕ'] = 'Командиру військової частини А5003'

    ctx['СЛІДЧИЙ_ПОСАДА'] = row.get('СЛІДЧИЙ_ПОСАДА', '')
    ctx['СЛІДЧИЙ_ПОСАДА_РОД'] = to_genitive(ctx['СЛІДЧИЙ_ПОСАДА'])
    ctx['СЛІДЧИЙ_ЗВАННЯ'] = str(row.get('СЛІДЧИЙ_ЗВАННЯ','') or '')
    ctx['СЛІДЧИЙ_ПІБ'] = format_name(str(row.get('СЛІДЧИЙ_ПІБ','') or ''))
    ctx['СЛІДЧИЙ_ЗВАННЯ_РОД'] = to_genitive(ctx['СЛІДЧИЙ_ЗВАННЯ'])
    ctx['СЛІДЧИЙ_ПІБ_РОД'] = format_name(to_genitive(ctx['СЛІДЧИЙ_ПІБ']))

    ctx['ПОДАВАЧ_ПОСАДА'] = str(row.get('ПОДАВАЧ_ПОСАДА', '') or '')
    ctx['ПОДАВАЧ_ПОСАДА_РОД'] = to_genitive(ctx['ПОДАВАЧ_ПОСАДА'])
    ctx['ПОДАВАЧ_ЗВАННЯ'] = str(row.get('ПОДАВАЧ_ЗВАННЯ','') or '')
    ctx['ПОДАВАЧ_ПІБ'] = str(row.get('ПОДАВАЧ_ПІБ','') or '')
    ctx['ПОДАВАЧ_ЗВАННЯ_РОД'] = to_genitive(ctx['ПОДАВАЧ_ЗВАННЯ'])
    ctx['ПОДАВАЧ_ПІБ_РОД'] = format_name(to_genitive(ctx['ПОДАВАЧ_ПІБ']))
    ctx['ПОДАВАЧ_ПІБ_КРАТКО'] = short_fio(ctx['ПОДАВАЧ_ПІБ'])

    ctx['ПІБ'] = format_name(fio)
    ctx['ЗВАННЯ'] = str(row.get('ЗВАННЯ','') or '')
    ctx['ЗВАННЯ_РОД'] = to_genitive(ctx['ЗВАННЯ'])
    ctx['ПІБ_РОД'] = format_name(to_genitive(fio))

    ctx['ТЕРРІТОРІЯ_ЦПП'] = str(row.get('ТЕРРІТОРІЯ_ЦПП','') or '')
    ctx['ТЕРРІТОРІЯ_НП'] = str(row.get('ТЕРРІТОРІЯ_НП','') or '')
    ctx['ТЕРРІТОРІЯ_РН'] = str(row.get('ТЕРРІТОРІЯ_РН','') or '')
    ctx['ДОПОВІДЬ_ВІД'] = str(row.get('ДОПОВІДЬ_ВІД','') or '')
    ctx['ПІДРОЗДІЛ'] = department
    ctx['ПІДРОЗДІЛ_РОД'] =  to_genitive(department)
    ctx['ЗБРОЯ'] = str(row.get('ЗБРОЯ','') or '')

    dopov_cherg = row.get('ЧАС_ДОПОВІДІ_ЧЕРГОВОМУ')
    tstr = str(dopov_cherg) if pd.notna(dopov_cherg) else ''

    ctx['ЧАС_ДОП_ЧЕРГ'] = ' о ' + tstr.split(':')[0] + ':' + tstr.split(':')[1] + ' год' if ':' in tstr else tstr

    date_val = row.get('СЗЧ_ДАТА')
    if pd.notna(date_val):
        dt = pd.to_datetime(date_val, dayfirst=True)
        ctx.update(date_parts(dt))
    else:
        ctx['СЗЧ_ДАТА'] = ''

    tval = row.get('СЗЧ_ЧАС')
    ctx['ТЕРМІН_РОЗСЛІДУВАННЯ'] = get_investigation_deadline(date_val, 11)
    tstr = str(tval) if pd.notna(tval) else ''
    ctx['СЗЧ_ЧАС'] = format_time_str(tval)

    obst = row.get('ОБСТАВИНИ')
    ctx['ОБСТАВИНИ'] = ', ' + str(obst).strip() if pd.notna(obst) and str(obst).strip() else ''

    dict_tamplates_doc = Document(resource_path('ШАБЛОНИ/СЛОВНИК ПІДШАБЛОНІВ.docx'))

    for var, col in VAR_TO_COL.items():
        raw = row.get(col)
        formatter = formatters.get(var, default_format)
        ctx[var] = formatter(raw)

    ctx['ДОПОВІДЬ_НОМЕР'] = ''

    suffix = fio.split()[0].upper() + ' ' + ''.join([p[0].upper()+'.' for p in fio.split()[1:]])

    # Генерація довідки
    if 'dovidka' in templates:
        doc = Document(templates['dovidka'])
        replace_placeholders(doc, ctx)

        if maino_value == "майно втрачене":
            maino_block = extract_block_by_tags(dict_tamplates_doc, "БЛОК_МАЙНО_ВТРАЧЕНЕ_ДОВІДКА")
        else:
            maino_block = extract_block_by_tags(dict_tamplates_doc, "БЛОК_МАЙНО_НАЯВНЕ_ДОВІДКА")
        
        insert_block_by_placeholder(doc, "{{МАЙНО}}", maino_block)

        filename = ("ДД_СЗЧ_" if format_value == "47" else "А5003 Довідка_") + f"{suffix}.docx"
        doc.save(os.path.join(OUTPUT_DIR, filename))

    # Генерація рапорту з блоком ПОДАВАЧ
    if 'raport' in templates:
        raport_template_path = templates['raport']
        doc = Document(raport_template_path)
        podavach_fio_raw = row.get('ПОДАВАЧ_ПІБ', '')
        ctx['ПОДАВАЧ'] = ''
        replace_placeholders(doc, ctx)
        
        if maino_value == "майно втрачене":
            maino_block = extract_block_by_tags(dict_tamplates_doc, "БЛОК_МАЙНО_ВТРАЧЕНЕ")
        else:
            maino_block = extract_block_by_tags(dict_tamplates_doc, "БЛОК_МАЙНО_НАЯВНЕ")

        insert_block_by_placeholder(doc, "{{МАЙНО}}", maino_block)

        szch_date = ctx.get('СЗЧ_ДАТА', '').strip()
        filename = f"А5003 РАПОРТ {fio.split()[0].upper()} {szch_date}.docx"
        doc.save(os.path.join(OUTPUT_DIR, filename))

print("Всього рядків у Excel:", len(df))
print('Готово! Документи збережено в', OUTPUT_DIR)
